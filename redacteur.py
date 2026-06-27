"""Agent rédacteur : assemble dossier.md, génère dossier.pdf, dossier.docx et prd.md."""
import datetime
import os
import re
from pathlib import Path

import io
import markdown
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, RGBColor
from openai import AsyncOpenAI

OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen3:14b")

# CSS simplifié pour xhtml2pdf (pas de @page avancé, pas de border-radius)
_CSS = """
body { font-family: Helvetica, Arial, sans-serif; font-size: 10pt; line-height: 1.6; color: #1a1a1a; }
h1 { font-size: 18pt; border-bottom: 2px solid #333; padding-bottom: 4px; margin-top: 0; }
h2 { font-size: 13pt; border-bottom: 1px solid #ccc; margin-top: 20px; color: #222; }
h3 { font-size: 11pt; margin-top: 14px; color: #333; }
blockquote { border-left: 4px solid #e0a000; margin: 10px 0; padding: 4px 12px;
             background: #fffbf0; font-style: italic; font-size: 9pt; }
table { border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 9pt; }
th { background: #eeeeee; border: 1px solid #cccccc; padding: 5px 8px; text-align: left; }
td { border: 1px solid #dddddd; padding: 4px 8px; }
a { color: #0055aa; }
code { background: #f5f5f5; padding: 1px 3px; font-size: 9pt; }
ul, ol { padding-left: 20px; }
li { margin-bottom: 3px; }
hr { border-top: 1px solid #dddddd; margin: 16px 0; }
"""

_GABARIT = """\
# Dossier — {idea}

> *Généré le {date} par le système multi-agents local (Ollama / {model}).*
> *Les informations ci-dessous sont issues de sources web publiques — vérifiez avant toute décision.*

---

## 1. Vue d'ensemble

{overview}

---

## 2. Skills & Connecteurs pertinents

{skills}

---

## 3. Analyse concurrentielle

{concurrents}

---

## 4. Marché & Faisabilité technique

{marche_techno}

---

## 5. Cadre juridique

> **⚠ Non-conseil juridique** — Cette section est fournie à titre informatif uniquement et \
ne constitue pas un avis juridique. Consultez un professionnel du droit pour toute décision.

{juridique}

---

## 6. Recommandations de stack

{stack}

---

*Dossier produit en local — 100 % privé, 0 € récurrent.*
"""


async def run_redacteur(
    llm: AsyncOpenAI,
    model: str,
    idea: str,
    overview: str,
    skills: str,
    concurrents: str,
    marche_techno: str,
    juridique: str,
    out_path: Path,
    log=print,
) -> tuple[Path, Path, Path, Path]:
    """Assemble dossier.md + .pdf + .docx + prd.md. Retourne les 4 chemins."""

    # 1. Recommandations de stack (appel LLM rapide)
    log("  → Génération recommandations de stack…")
    stack = await _gen_stack(llm, model, idea, skills, concurrents)

    # 2. Assembler dossier.md
    date = datetime.date.today().isoformat()
    dossier_md = _GABARIT.format(
        idea=idea, date=date, model=model,
        overview=overview, skills=skills,
        concurrents=concurrents, marche_techno=marche_techno,
        juridique=juridique, stack=stack,
    )
    md_path = out_path / "dossier.md"
    md_path.write_text(dossier_md, encoding="utf-8")
    log(f"  → dossier.md écrit ({len(dossier_md)} chars)")

    # 3. dossier.pdf via xhtml2pdf
    pdf_path = out_path / "dossier.pdf"
    _render_pdf(dossier_md, pdf_path)
    log("  → dossier.pdf généré")

    # 4. dossier.docx via python-docx
    docx_path = out_path / "dossier.docx"
    _render_docx(dossier_md, docx_path)
    log("  → dossier.docx généré")

    # 5. prd.md
    prd_path = out_path / "prd.md"
    prd_md = await _gen_prd(llm, model, idea, dossier_md)
    prd_path.write_text(prd_md, encoding="utf-8")
    log("  → prd.md écrit")

    return md_path, pdf_path, docx_path, prd_path


async def _gen_stack(llm, model, idea, skills, concurrents) -> str:
    resp = await llm.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": (
                "Tu es un architecte technique. À partir des outils/connecteurs disponibles et de l'analyse concurrentielle, "
                "recommande une stack technique concrète pour construire ce projet (backend, frontend, BDD, hébergement, IA). "
                "Format : bullet points avec justification courte. Cite tes sources si tu en as. Réponds en français."
            )},
            {"role": "user", "content": (
                f"Projet : {idea}\n\n"
                f"## Outils & connecteurs disponibles\n{skills}\n\n"
                f"## Analyse concurrentielle\n{concurrents}"
            )},
        ],
        temperature=0.3,
        extra_body={"think": False},
    )
    return resp.choices[0].message.content.strip()


async def _gen_prd(llm, model, idea, dossier_md) -> str:
    resp = await llm.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": (
                "Tu es un product manager. À partir du dossier fourni, rédige un PRD (Product Requirements Document) structuré :\n"
                "1. Vision & objectifs (3-5 bullets)\n"
                "2. Utilisateurs cibles & besoins\n"
                "3. Fonctionnalités Must/Should/Could (tableau MoSCoW)\n"
                "4. Stack recommandée\n"
                "5. Risques & mitigations\n"
                "6. Definition of Done (checklist)\n"
                "Sois précis, factuel, cite les sources du dossier. Réponds en français.\n"
                "Ajoute l'avertissement 'non-conseil juridique' si le dossier contient une section juridique."
            )},
            {"role": "user", "content": f"Dossier :\n\n{dossier_md}"},
        ],
        temperature=0.3,
        extra_body={"think": False},
    )
    return resp.choices[0].message.content.strip()


def _render_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in md_text.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles["Quote"] if "Quote" in doc.styles else doc.styles["Normal"]
            p.runs[0].italic = True
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        elif stripped == "---":
            doc.add_paragraph("─" * 60)
        elif stripped == "":
            doc.add_paragraph("")
        else:
            # Nettoyer le markdown inline (**bold**, *italic*, `code`)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*",     r"\1", clean)
            clean = re.sub(r"`(.+?)`",        r"\1", clean)
            doc.add_paragraph(clean)

    doc.save(out_path)


def _render_pdf(md_text: str, out_path: Path) -> None:
    html_body = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "nl2br"],
    )
    html = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<style>{_CSS}</style></head><body>{html_body}</body></html>"
    )
    buf = io.BytesIO()
    result = pisa.CreatePDF(io.StringIO(html), dest=buf)
    if result.err:
        raise RuntimeError(f"xhtml2pdf error : {result.err}")
    out_path.write_bytes(buf.getvalue())
